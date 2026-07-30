"""Tests for novels/importer.py — markdown/docx parsing and file dispatch.

Usage:
    cd client/backend
    python -m pytest tests/test_importer.py -v
"""

import os
import tempfile

import pytest

from novels.importer import (
    ChapterData,
    VolumeData,
    ParseResult,
    ParseWarning,
    parse_docx,
    parse_file,
    parse_markdown,
)

FIXTURE_DIR = os.path.join(os.path.dirname(__file__), "import")


# =========================================================================
#  parse_markdown
# =========================================================================


class TestParseMarkdown:
    def test_valid_2vol_3ch(self):
        """Standard 2 volumes, 3 chapters with full structure."""
        path = os.path.join(FIXTURE_DIR, "valid-2vol-3ch.txt")
        with open(path, encoding="utf-8") as f:
            content = f.read()

        result = parse_markdown(content, "valid-2vol-3ch.txt")

        assert len(result.volumes) == 2
        assert result.volumes[0].title == "第一卷：风云初起"
        assert result.volumes[1].title == "第二卷：暗流涌动"

        vol1_chs = result.volumes[0].chapters
        assert len(vol1_chs) == 2
        assert vol1_chs[0].title == "第一章：初入江湖"
        assert vol1_chs[1].title == "第二章：山门试炼"

        vol2_chs = result.volumes[1].chapters
        assert len(vol2_chs) == 1
        assert vol2_chs[0].title == "第三章：宗门疑云"

        # Word counts should be > 0
        assert vol1_chs[0].word_count > 0
        assert vol1_chs[1].word_count > 0
        assert vol2_chs[0].word_count > 0

        # Title extracted from first # heading
        assert result.title == "第一卷：风云初起"

        # No warnings for clean input
        assert len(result.warnings) == 0

    def test_single_chapter_fallback_volume(self):
        """Only ## headings — auto-create default volume '正文'."""
        path = os.path.join(FIXTURE_DIR, "single-chapter-only.md")
        with open(path, encoding="utf-8") as f:
            content = f.read()

        result = parse_markdown(content, "single-chapter-only.md")

        assert len(result.volumes) == 1
        assert result.volumes[0].title == "正文"
        assert len(result.volumes[0].chapters) == 1
        assert result.volumes[0].chapters[0].title == "唯一的一章正文"
        assert result.volumes[0].chapters[0].word_count > 0

    def test_orphan_text_warning(self):
        """Text before first heading triggers orphan_text warning."""
        path = os.path.join(FIXTURE_DIR, "orphan-text.txt")
        with open(path, encoding="utf-8") as f:
            content = f.read()

        result = parse_markdown(content, "orphan-text.txt")

        # Should still parse the volume and chapter
        assert len(result.volumes) == 1
        assert len(result.volumes[0].chapters) >= 1

        # Should have orphan_text warning
        warning_types = [w.type for w in result.warnings]
        assert "orphan_text" in warning_types
        orphan_warnings = [w for w in result.warnings if w.type == "orphan_text"]
        assert len(orphan_warnings) >= 1
        assert orphan_warnings[0].details is not None
        assert "preview" in orphan_warnings[0].details

    def test_no_structure(self):
        """Plain text with no # or ## headings."""
        content = "这是纯文本内容\n没有任何标题结构\n只有几行文字。\n"

        result = parse_markdown(content, "plain.txt")

        assert len(result.volumes) == 0
        assert len(result.warnings) >= 1
        warning_types = [w.type for w in result.warnings]
        assert "no_volume_title" in warning_types
        assert result.title == "plain"

    def test_empty_file(self):
        """Empty string produces empty volumes with warning."""
        result = parse_markdown("", "empty.txt")

        assert len(result.volumes) == 0
        assert len(result.warnings) >= 1
        assert result.warnings[0].type == "no_volume_title"

    def test_title_from_file_when_no_volume(self):
        """When no # heading, title falls back to filename."""
        content = "## 第一章\n一些正文内容。\n"
        result = parse_markdown(content, "my-novel.txt")

        assert result.title == "my-novel"

    def test_title_unnamed_fallback(self):
        """When no structure and no filename, default to '未命名作品'."""
        result = parse_markdown("随便写写", "")

        assert result.title == "未命名作品"

    def test_body_text_in_volume_no_chapter_auto_creates(self):
        """Text inside a volume but before any ## auto-creates a default chapter."""
        content = "# 第一卷\n这是卷内的正文，没有章节标题。\n## 第二章\n这是第二章。\n"
        result = parse_markdown(content)

        assert len(result.volumes) == 1
        assert len(result.volumes[0].chapters) == 2
        assert result.volumes[0].chapters[0].title == "第1章"
        assert result.volumes[0].chapters[0].content.strip() == "这是卷内的正文，没有章节标题。"
        assert result.volumes[0].chapters[1].title == "第二章"

    def test_multiple_orphan_blocks_give_single_warning(self):
        """Multiple orphan text lines all get collected into one warning."""
        content = "游离行1\n游离行2\n\n游离行3\n# 卷\n## 章\n正文\n"
        result = parse_markdown(content)

        orphan_warnings = [w for w in result.warnings if w.type == "orphan_text"]
        assert len(orphan_warnings) == 1
        assert "3" in orphan_warnings[0].message


# =========================================================================
#  parse_docx
# =========================================================================


class TestParseDocx:
    def test_valid(self):
        """Parse a docx file with Heading 1 and Heading 2 structure."""
        docx_path = os.path.join(FIXTURE_DIR, "valid-2vol-3ch.docx")
        if not os.path.exists(docx_path):
            # Create a minimal docx fixture
            _create_test_docx(docx_path)

        result = parse_docx(docx_path)

        assert len(result.volumes) >= 2
        assert result.volumes[0].title == "第一卷：风云初起"
        # Clean up
        os.unlink(docx_path)

    def test_missing_heading(self):
        """Document with no Heading 1 or Heading 2 produces warnings."""
        docx_path = os.path.join(FIXTURE_DIR, "no-heading.docx")
        if not os.path.exists(docx_path):
            _create_plain_docx(docx_path)

        result = parse_docx(docx_path)

        assert len(result.volumes) == 0
        assert any(w.type == "no_volume_title" for w in result.warnings)
        os.unlink(docx_path)


# =========================================================================
#  parse_file (dispatcher)
# =========================================================================


class TestParseFile:
    def test_md_extension_dispatches_to_markdown(self):
        """.md file dispatches to parse_markdown."""
        with tempfile.NamedTemporaryFile(
            suffix=".md", mode="w", encoding="utf-8", delete=False
        ) as f:
            f.write("# 卷\n## 章\n正文\n")
            tmp = f.name

        try:
            result = parse_file(tmp, "test.md")
            assert len(result.volumes) == 1
            assert result.volumes[0].title == "卷"
        finally:
            os.unlink(tmp)

    def test_txt_extension_dispatches_to_markdown(self):
        """.txt file dispatches to parse_markdown."""
        with tempfile.NamedTemporaryFile(
            suffix=".txt", mode="w", encoding="utf-8", delete=False
        ) as f:
            f.write("## 章\n正文\n")
            tmp = f.name

        try:
            result = parse_file(tmp, "test.txt")
            assert len(result.volumes) == 1
        finally:
            os.unlink(tmp)

    def test_docx_extension_dispatches_to_docx(self):
        """.docx file dispatches to parse_docx."""
        docx_path = os.path.join(FIXTURE_DIR, "_dispatch_test.docx")
        if not os.path.exists(docx_path):
            _create_test_docx(docx_path)

        try:
            result = parse_file(docx_path, "_dispatch_test.docx")
            assert len(result.volumes) >= 2
        finally:
            if os.path.exists(docx_path):
                os.unlink(docx_path)

    def test_invalid_extension_raises(self):
        """Unsupported extension raises ValueError."""
        with pytest.raises(ValueError, match="不支持的文件格式"):
            # File path doesn't even need to exist — dispatcher checks extension
            parse_file("/nonexistent/test.pdf", "test.pdf")

    def test_empty_filename(self):
        """Empty filename falls back to filepath extension detection."""
        with tempfile.NamedTemporaryFile(
            suffix=".txt", mode="w", encoding="utf-8", delete=False
        ) as f:
            f.write("内容")
            tmp = f.name

        try:
            # parse_file falls back to filepath's .txt extension
            result = parse_file(tmp, "")
            # No # heading and no filename → default title
            assert result.title == "未命名作品"
            assert len(result.warnings) >= 1
        finally:
            os.unlink(tmp)


# =========================================================================
#  Helpers
# =========================================================================


def _create_test_docx(path: str):
    """Create a minimal .docx fixture with Heading 1 and Heading 2 styles."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Heading 1 — volume
    doc.add_heading("第一卷：风云初起", level=1)
    doc.add_paragraph("这是第一卷的简介段落。")

    # Heading 2 — chapter
    doc.add_heading("第一章：初入江湖", level=2)
    doc.add_paragraph("少年林逸站在青云山下。")
    doc.add_paragraph("他迈出了第一步。")

    # Heading 2 — another chapter
    doc.add_heading("第二章：山门试炼", level=2)
    doc.add_paragraph("青云宗的山门前，少年少女正在排队。")

    # Heading 1 — another volume
    doc.add_heading("第二卷：暗流涌动", level=1)

    doc.add_heading("第三章：宗门疑云", level=2)
    doc.add_paragraph("林逸发现宗门内部暗流涌动。")

    doc.save(path)


def _create_plain_docx(path: str):
    """Create a .docx with no heading styles (plain paragraphs only)."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("这是一段普通文字。")
    doc.add_paragraph("没有任何标题样式。")
    doc.add_paragraph("只有几段文字。")

    doc.save(path)
